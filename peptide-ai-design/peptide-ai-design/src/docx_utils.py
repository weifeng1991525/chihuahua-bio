#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docx_utils.py - DOCX格式工具
定义文档格式常量和辅助函数
格式要求：中文宋体小四，英文Times New Roman小四，行间距1.5倍，段前段后0，首行缩进2字符
"""

from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os


# ==================== 格式常量 ====================

# 字体
FONT_CN = '宋体'
FONT_EN = 'Times New Roman'
FONT_TITLE_CN = '黑体'
FONT_TITLE_EN = 'Arial'

# 字号
FONT_SIZE_XIAOSI = Pt(12)  # 小四 = 12pt
FONT_SIZE_SANHAO = Pt(16)  # 三号 = 16pt
FONT_SIZE_XIAOSAN = Pt(15)  # 小三 = 15pt
FONT_SIZE_ERHAO = Pt(22)   # 二号 = 22pt
FONT_SIZE_XIAOER = Pt(18)  # 小二 = 18pt
FONT_SIZE_SIHAO = Pt(14)   # 四号 = 14pt
FONT_SIZE_WUHAO = Pt(10.5) # 五号 = 10.5pt
FONT_SIZE_XIAOWU = Pt(9)   # 小五 = 9pt

# 行间距
LINE_SPACING_15 = 1.5

# 页边距
MARGIN_TOP = Cm(2.54)
MARGIN_BOTTOM = Cm(2.54)
MARGIN_LEFT = Cm(3.17)
MARGIN_RIGHT = Cm(3.17)

# 颜色
COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_DARK_GRAY = RGBColor(64, 64, 64)
COLOR_GRAY = RGBColor(128, 128, 128)
COLOR_BLUE = RGBColor(0, 51, 153)
COLOR_RED = RGBColor(192, 0, 0)
COLOR_TABLE_HEADER_BG = RGBColor(0, 51, 102)
COLOR_TABLE_HEADER_TEXT = RGBColor(255, 255, 255)
COLOR_TABLE_ALT_BG = RGBColor(240, 245, 250)


# ==================== 段落格式 ====================

def set_paragraph_format(paragraph, font_size=FONT_SIZE_XIAOSI, bold=False,
                         alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=None, space_before=0, space_after=0,
                         line_spacing=LINE_SPACING_15, font_cn=FONT_CN, font_en=FONT_EN):
    """设置段落格式"""
    pf = paragraph.paragraph_format
    pf.alignment = alignment
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing

    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        # 首行缩进2字符 = 小四字号 * 2
        pf.first_line_indent = Cm(0.74)  # 约2个中文字符

    # 设置字体
    for run in paragraph.runs:
        run.font.size = font_size
        run.font.bold = bold
        run.font.name = font_en
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
        run.font.color.rgb = COLOR_BLACK

    return paragraph


def add_formatted_paragraph(doc, text, font_size=FONT_SIZE_XIAOSI, bold=False,
                           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           first_line_indent=True, space_before=0, space_after=0,
                           line_spacing=LINE_SPACING_15, font_cn=FONT_CN, font_en=FONT_EN,
                           color=None):
    """添加格式化段落（支持中英文混排字体）"""
    paragraph = doc.add_paragraph()
    pf = paragraph.paragraph_format
    pf.alignment = alignment
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing

    if first_line_indent and alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        pf.first_line_indent = Cm(0.74)

    # 智能分词：中文用中文字体，英文/数字用英文字体
    if text:
        run = paragraph.add_run(text)
        run.font.size = font_size
        run.font.bold = bold
        run.font.name = font_en
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
        if color:
            run.font.color.rgb = color

    return paragraph


def add_mixed_paragraph(doc, segments, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                        first_line_indent=True, space_before=0, space_after=0,
                        line_spacing=LINE_SPACING_15):
    """
    添加中英文混排段落
    segments: list of dict, 每个dict包含:
        - text: 文本内容
        - bold: 是否加粗
        - font_cn: 中文字体
        - font_en: 英文字体
        - font_size: 字号
        - color: 颜色
        - italic: 斜体
    """
    paragraph = doc.add_paragraph()
    pf = paragraph.paragraph_format
    pf.alignment = alignment
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing

    if first_line_indent and alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        pf.first_line_indent = Cm(0.74)

    for seg in segments:
        text = seg.get('text', '')
        if not text:
            continue
        run = paragraph.add_run(text)
        run.font.size = seg.get('font_size', FONT_SIZE_XIAOSI)
        run.font.bold = seg.get('bold', False)
        run.font.italic = seg.get('italic', False)
        run.font.name = seg.get('font_en', FONT_EN)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), seg.get('font_cn', FONT_CN))
        if seg.get('color'):
            run.font.color.rgb = seg['color']

    return paragraph


def add_heading_cn(doc, text, level=1):
    """添加中文标题"""
    if level == 1:
        return add_formatted_paragraph(doc, text, font_size=FONT_SIZE_SANHAO, bold=True,
                                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                       first_line_indent=False, space_before=12, space_after=6,
                                       font_cn=FONT_TITLE_CN, font_en=FONT_TITLE_EN)
    elif level == 2:
        return add_formatted_paragraph(doc, text, font_size=FONT_SIZE_SIHAO, bold=True,
                                       alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                       first_line_indent=False, space_before=8, space_after=4,
                                       font_cn=FONT_TITLE_CN, font_en=FONT_TITLE_EN)
    elif level == 3:
        return add_formatted_paragraph(doc, text, font_size=FONT_SIZE_XIAOSI, bold=True,
                                       alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                       first_line_indent=False, space_before=6, space_after=3,
                                       font_cn=FONT_TITLE_CN, font_en=FONT_TITLE_EN)
    else:
        return add_formatted_paragraph(doc, text, font_size=FONT_SIZE_XIAOSI, bold=True,
                                       alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                       first_line_indent=False, space_before=4, space_after=2,
                                       font_cn=FONT_TITLE_CN, font_en=FONT_TITLE_EN)


# ==================== 表格格式 ====================

def set_cell_format(cell, text, font_size=FONT_SIZE_XIAOWU, bold=False,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    font_cn=FONT_CN, font_en=FONT_EN,
                    bg_color=None, text_color=None):
    """设置单元格格式"""
    cell.text = ''
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.0

    run = paragraph.add_run(str(text))
    run.font.size = font_size
    run.font.bold = bold
    run.font.name = font_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
    if text_color:
        run.font.color.rgb = text_color

    if bg_color:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # 垂直居中
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcVAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(tcVAlign)

    return cell


def create_table(doc, headers, rows, col_widths=None):
    """创建格式化表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 表头
    header_color = "003366"
    for j, header in enumerate(headers):
        set_cell_format(table.rows[0].cells[j], header,
                       font_size=FONT_SIZE_XIAOWU, bold=True,
                       bg_color=header_color,
                       text_color=COLOR_TABLE_HEADER_TEXT)

    # 数据行
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            bg = "F0F5FA" if i % 2 == 0 else None
            set_cell_format(table.rows[i + 1].cells[j], str(cell_text),
                           font_size=FONT_SIZE_XIAOWU,
                           bg_color=bg)

    # 设置列宽
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_table_caption(doc, caption, table_num="Table 1"):
    """添加表格标题"""
    add_formatted_paragraph(doc, f'{table_num} {caption}',
                           font_size=FONT_SIZE_WUHAO, bold=True,
                           alignment=WD_ALIGN_PARAGRAPH.CENTER,
                           first_line_indent=False, space_before=6, space_after=2)


# ==================== 图片格式 ====================

def add_figure(doc, image_path, caption="", figure_num="Fig.1", width=Cm(15)):
    """添加图片和图注"""
    if not os.path.exists(image_path):
        add_formatted_paragraph(doc, f'[图片缺失: {image_path}]',
                               font_size=FONT_SIZE_WUHAO, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                               first_line_indent=False, color=COLOR_RED)
        return

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)

    run = paragraph.add_run()
    run.add_picture(image_path, width=width)

    if caption:
        add_formatted_paragraph(doc, f'{figure_num} {caption}',
                               font_size=FONT_SIZE_WUHAO, bold=False,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER,
                               first_line_indent=False, space_before=2, space_after=6)


# ==================== 页面设置 ====================

def set_page_format(doc):
    """设置页面格式"""
    section = doc.sections[0]
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)

    # 添加页眉页脚
    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("多肽AI辅助设计方案")
    run.font.size = FONT_SIZE_WUHAO
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)
    run.font.color.rgb = COLOR_GRAY

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("第 ")
    run.font.size = FONT_SIZE_WUHAO
    run.font.name = FONT_EN
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)

    # 添加页码字段
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run._element.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run._element.append(fldChar2)

    run2 = footer_para.add_run(" 页")
    run2.font.size = FONT_SIZE_WUHAO
    run2.font.name = FONT_EN
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)


def add_page_break(doc):
    """添加分页符"""
    doc.add_page_break()


def add_horizontal_line(doc):
    """添加水平分割线"""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="003366"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
