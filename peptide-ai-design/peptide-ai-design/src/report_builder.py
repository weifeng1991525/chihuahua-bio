#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
report_builder.py - 多肽设计方案DOCX报告构建器
基于分析结果构建完整的多肽设计方案文档
"""

import os
import json
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from docx_utils import (
    set_page_format, add_formatted_paragraph, add_mixed_paragraph,
    add_heading_cn, create_table, add_table_caption,
    add_figure, add_page_break, add_horizontal_line,
    FONT_CN, FONT_EN, FONT_TITLE_CN, FONT_TITLE_EN,
    FONT_SIZE_XIAOSI, FONT_SIZE_SANHAO, FONT_SIZE_SIHAO,
    FONT_SIZE_XIAOER, FONT_SIZE_ERHAO, FONT_SIZE_WUHAO, FONT_SIZE_XIAOWU,
    COLOR_BLACK, COLOR_DARK_GRAY, COLOR_GRAY, COLOR_BLUE, COLOR_RED,
    LINE_SPACING_15
)


class PeptideReportBuilder:
    """多肽设计方案报告构建器"""

    def __init__(self):
        self.doc = Document()
        self._setup_document()

    def _setup_document(self):
        """初始化文档设置"""
        set_page_format(self.doc)
        # 设置默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = FONT_EN
        font.size = FONT_SIZE_XIAOSI
        style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)

    def build(self, report_data: Dict[str, Any], output_path: str) -> str:
        """
        构建完整报告

        Args:
            report_data: 完整的报告JSON数据
            output_path: 输出文件路径

        Returns:
            输出文件路径
        """
        print("  [报告构建] 开始构建多肽设计方案DOCX报告...")

        # 1. 封面
        self._build_cover(report_data)

        # 2. 项目概述
        self._build_overview(report_data)

        # 3. 需求分析
        self._build_requirement_analysis(report_data)

        # 4. 多肽设计方案
        self._build_peptide_design(report_data)

        # 5. 合成方案
        self._build_synthesis_plan(report_data)

        # 6. 修饰方案
        self._build_modification_plan(report_data)

        # 7. 质量控制
        self._build_quality_control(report_data)

        # 8. 项目周期与报价
        self._build_timeline_and_pricing(report_data)

        # 9. 风险评估与建议
        self._build_risk_assessment(report_data)

        # 10. 附录
        self._build_appendix(report_data)

        # 保存
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
        print(f"  [报告构建] 报告已保存: {output_path}")
        return output_path

    # ==================== 封面 ====================

    def _build_cover(self, data: Dict[str, Any]):
        """构建封面"""
        project_info = data.get("project_info", {})
        title = project_info.get("title", "多肽AI辅助设计方案")
        project_id = project_info.get("project_id", "")
        date = project_info.get("date", "")
        client = project_info.get("client", "")

        # 空行占位
        for _ in range(4):
            self.doc.add_paragraph()

        # 机构名称
        add_formatted_paragraph(self.doc, "多肽AI辅助设计中心",
                               font_size=FONT_SIZE_XIAOER, bold=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER,
                               first_line_indent=False, space_before=0, space_after=20,
                               font_cn=FONT_TITLE_CN, font_en=FONT_TITLE_EN,
                               color=RGBColor(0, 51, 102))

        # 分割线
        add_horizontal_line(self.doc)

        # 空行
        for _ in range(2):
            self.doc.add_paragraph()

        # 报告标题
        add_formatted_paragraph(self.doc, title,
                               font_size=FONT_SIZE_ERHAO, bold=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER,
                               first_line_indent=False, space_before=20, space_after=10,
                               font_cn=FONT_TITLE_CN, font_en=FONT_TITLE_EN,
                               color=RGBColor(0, 51, 102))

        # 英文标题
        add_formatted_paragraph(self.doc, project_info.get("title_en", "Peptide AI-Assisted Design Proposal"),
                               font_size=FONT_SIZE_SIHAO, bold=False,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER,
                               first_line_indent=False, space_before=0, space_after=20,
                               font_cn=FONT_CN, font_en=FONT_EN,
                               color=COLOR_DARK_GRAY)

        # 分割线
        add_horizontal_line(self.doc)

        # 空行
        for _ in range(3):
            self.doc.add_paragraph()

        # 项目信息表格
        info_items = [
            ("项目编号", project_id),
            ("编制日期", date),
            ("客户名称", client),
            ("方案版本", project_info.get("version", "V1.0")),
            ("保密级别", project_info.get("confidentiality", "机密")),
        ]

        for label, value in info_items:
            add_mixed_paragraph(self.doc, [
                {"text": f"{label}：", "bold": True, "font_size": FONT_SIZE_SIHAO,
                 "font_cn": FONT_TITLE_CN, "font_en": FONT_TITLE_EN,
                 "color": RGBColor(0, 51, 102)},
                {"text": value, "bold": False, "font_size": FONT_SIZE_SIHAO,
                 "font_cn": FONT_CN, "font_en": FONT_EN}
            ], alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=False,
               space_before=4, space_after=4)

        # 空行
        for _ in range(4):
            self.doc.add_paragraph()

        # 机构落款
        add_formatted_paragraph(self.doc, "多肽AI辅助设计中心",
                               font_size=FONT_SIZE_SIHAO, bold=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER,
                               first_line_indent=False,
                               font_cn=FONT_TITLE_CN, font_en=FONT_TITLE_EN)

        add_page_break(self.doc)

    # ==================== 项目概述 ====================

    def _build_overview(self, data: Dict[str, Any]):
        """构建项目概述"""
        add_heading_cn(self.doc, "一、项目概述", level=1)
        add_horizontal_line(self.doc)

        overview = data.get("overview", {})
        
        # 项目背景
        add_heading_cn(self.doc, "1.1 项目背景", level=2)
        background = overview.get("background", "")
        if background:
            add_formatted_paragraph(self.doc, background,
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=6, space_after=6)

        # 设计目标
        add_heading_cn(self.doc, "1.2 设计目标", level=2)
        objectives = overview.get("objectives", [])
        for obj in objectives:
            add_formatted_paragraph(self.doc, f"● {obj}",
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=2, space_after=2)

        # 设计原则
        add_heading_cn(self.doc, "1.3 设计原则", level=2)
        principles = overview.get("principles", [])
        for principle in principles:
            add_formatted_paragraph(self.doc, f"● {principle}",
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=2, space_after=2)

        add_page_break(self.doc)

    # ==================== 需求分析 ====================

    def _build_requirement_analysis(self, data: Dict[str, Any]):
        """构建需求分析"""
        add_heading_cn(self.doc, "二、需求分析", level=1)
        add_horizontal_line(self.doc)

        requirements = data.get("requirements", {})

        # 功能需求
        add_heading_cn(self.doc, "2.1 功能需求", level=2)
        functional = requirements.get("functional", [])
        for req in functional:
            add_formatted_paragraph(self.doc, f"● {req}",
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=2, space_after=2)

        # 技术需求
        add_heading_cn(self.doc, "2.2 技术需求", level=2)
        technical = requirements.get("technical", [])
        for req in technical:
            add_formatted_paragraph(self.doc, f"● {req}",
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=2, space_after=2)

        # 约束条件
        add_heading_cn(self.doc, "2.3 约束条件", level=2)
        constraints = requirements.get("constraints", [])
        for constraint in constraints:
            add_formatted_paragraph(self.doc, f"● {constraint}",
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=2, space_after=2)

        # 需求汇总表
        add_heading_cn(self.doc, "2.4 需求汇总", level=2)
        summary_headers = ["需求类别", "需求内容", "优先级", "备注"]
        summary_rows = requirements.get("summary_table", [])
        if summary_rows:
            create_table(self.doc, summary_headers, summary_rows, col_widths=[3, 8, 2, 3])

        add_page_break(self.doc)

    # ==================== 多肽设计方案 ====================

    def _build_peptide_design(self, data: Dict[str, Any]):
        """构建多肽设计方案"""
        add_heading_cn(self.doc, "三、多肽设计方案", level=1)
        add_horizontal_line(self.doc)

        designs = data.get("peptide_designs", [])
        
        for idx, design in enumerate(designs):
            design_num = idx + 1
            add_heading_cn(self.doc, f"3.{design_num} 设计方案{design_num}", level=2)

            # 基本信息
            name = design.get("name", "")
            sequence = design.get("sequence", "")
            length = design.get("length", 0)
            mw = design.get("molecular_weight", 0)
            pi = design.get("pi", 0)

            add_mixed_paragraph(self.doc, [
                {"text": f"多肽名称：", "bold": True, "font_size": FONT_SIZE_XIAOSI,
                 "font_cn": FONT_TITLE_CN, "font_en": FONT_TITLE_EN},
                {"text": name, "font_size": FONT_SIZE_XIAOSI}
            ], first_line_indent=False, space_before=4, space_after=2)

            add_mixed_paragraph(self.doc, [
                {"text": f"氨基酸序列：", "bold": True, "font_size": FONT_SIZE_XIAOSI,
                 "font_cn": FONT_TITLE_CN, "font_en": FONT_TITLE_EN},
                {"text": sequence, "font_size": FONT_SIZE_XIAOSI, "font_en": FONT_EN}
            ], first_line_indent=False, space_before=2, space_after=2)

            # 理化参数表
            param_headers = ["参数", "数值", "说明"]
            param_rows = [
                ["序列长度", f"{length} aa", "氨基酸残基数"],
                ["分子量", f"{mw:.2f} Da", "理论分子量"],
                ["等电点", f"{pi:.2f}", "理论等电点"],
                ["净电荷(pH 7.0)", f"{design.get('net_charge', 0):+.0f}", "生理pH下的电荷"],
                ["疏水性", f"{design.get('hydrophobicity', 0):.2f}", "平均疏水指数"],
                ["溶解度建议", design.get("solubility_advice", ""), "溶解条件推荐"],
            ]
            create_table(self.doc, param_headers, param_rows, col_widths=[4, 4, 8])

            # 设计 rationale
            rationale = design.get("design_rationale", "")
            if rationale:
                add_heading_cn(self.doc, f"3.{design_num}.1 设计原理", level=3)
                add_formatted_paragraph(self.doc, rationale,
                                       font_size=FONT_SIZE_XIAOSI,
                                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                       first_line_indent=True, space_before=4, space_after=4)

            # 结构特征
            structure = design.get("structure_features", [])
            if structure:
                add_heading_cn(self.doc, f"3.{design_num}.2 结构特征", level=3)
                for feature in structure:
                    add_formatted_paragraph(self.doc, f"● {feature}",
                                           font_size=FONT_SIZE_XIAOSI,
                                           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                           first_line_indent=True, space_before=1, space_after=1)

            # 预期性质
            properties = design.get("expected_properties", [])
            if properties:
                add_heading_cn(self.doc, f"3.{design_num}.3 预期性质", level=3)
                for prop in properties:
                    add_formatted_paragraph(self.doc, f"● {prop}",
                                           font_size=FONT_SIZE_XIAOSI,
                                           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                           first_line_indent=True, space_before=1, space_after=1)

            add_page_break(self.doc)

    # ==================== 合成方案 ====================

    def _build_synthesis_plan(self, data: Dict[str, Any]):
        """构建合成方案"""
        add_heading_cn(self.doc, "四、合成方案", level=1)
        add_horizontal_line(self.doc)

        synthesis = data.get("synthesis", {})

        # 合成策略
        add_heading_cn(self.doc, "4.1 合成策略", level=2)
        strategy = synthesis.get("strategy", "")
        if strategy:
            add_formatted_paragraph(self.doc, strategy,
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=6, space_after=6)

        # 合成方法选择
        add_heading_cn(self.doc, "4.2 合成方法", level=2)
        method = synthesis.get("method", {})
        method_name = method.get("name", "")
        method_desc = method.get("description", "")
        
        add_formatted_paragraph(self.doc, f"推荐方法：{method_name}",
                               font_size=FONT_SIZE_XIAOSI, bold=True,
                               alignment=WD_ALIGN_PARAGRAPH.LEFT,
                               first_line_indent=False, space_before=4, space_after=2)
        
        if method_desc:
            add_formatted_paragraph(self.doc, method_desc,
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=2, space_after=4)

        # 合成参数
        add_heading_cn(self.doc, "4.3 合成参数", level=2)
        params = synthesis.get("parameters", {})
        param_rows = []
        for key, value in params.items():
            param_rows.append([key, str(value)])
        
        if param_rows:
            create_table(self.doc, ["参数", "设置"], param_rows, col_widths=[6, 10])

        # 合成步骤
        add_heading_cn(self.doc, "4.4 合成步骤", level=2)
        steps = synthesis.get("steps", [])
        for step in steps:
            step_num = step.get("step_number", "")
            step_title = step.get("title", "")
            step_desc = step.get("description", "")
            
            add_heading_cn(self.doc, f"步骤 {step_num}: {step_title}", level=3)
            if step_desc:
                add_formatted_paragraph(self.doc, step_desc,
                                       font_size=FONT_SIZE_XIAOSI,
                                       alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                       first_line_indent=True, space_before=2, space_after=4)

        # 纯化方案
        add_heading_cn(self.doc, "4.5 纯化方案", level=2)
        purification = synthesis.get("purification", {})
        purification_method = purification.get("method", "")
        if purification_method:
            add_formatted_paragraph(self.doc, f"纯化方法：{purification_method}",
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=4, space_after=4)
        
        purification_details = purification.get("details", [])
        for detail in purification_details:
            add_formatted_paragraph(self.doc, f"● {detail}",
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=1, space_after=1)

        add_page_break(self.doc)

    # ==================== 修饰方案 ====================

    def _build_modification_plan(self, data: Dict[str, Any]):
        """构建修饰方案"""
        add_heading_cn(self.doc, "五、修饰方案", level=1)
        add_horizontal_line(self.doc)

        modifications = data.get("modifications", [])
        
        if not modifications:
            add_formatted_paragraph(self.doc, "本方案不涉及化学修饰。",
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=6, space_after=6)
            add_page_break(self.doc)
            return

        for idx, mod in enumerate(modifications):
            mod_num = idx + 1
            add_heading_cn(self.doc, f"5.{mod_num} 修饰{mod_num}", level=2)

            mod_name = mod.get("name", "")
            mod_type = mod.get("type", "")
            mod_position = mod.get("position", "")
            mod_purpose = mod.get("purpose", "")

            add_mixed_paragraph(self.doc, [
                {"text": f"修饰名称：", "bold": True, "font_size": FONT_SIZE_XIAOSI},
                {"text": mod_name, "font_size": FONT_SIZE_XIAOSI}
            ], first_line_indent=False, space_before=4, space_after=2)

            add_mixed_paragraph(self.doc, [
                {"text": f"修饰类型：", "bold": True, "font_size": FONT_SIZE_XIAOSI},
                {"text": mod_type, "font_size": FONT_SIZE_XIAOSI}
            ], first_line_indent=False, space_before=2, space_after=2)

            add_mixed_paragraph(self.doc, [
                {"text": f"修饰位置：", "bold": True, "font_size": FONT_SIZE_XIAOSI},
                {"text": mod_position, "font_size": FONT_SIZE_XIAOSI}
            ], first_line_indent=False, space_before=2, space_after=2)

            add_mixed_paragraph(self.doc, [
                {"text": f"修饰目的：", "bold": True, "font_size": FONT_SIZE_XIAOSI},
                {"text": mod_purpose, "font_size": FONT_SIZE_XIAOSI}
            ], first_line_indent=False, space_before=2, space_after=4)

            # 修饰参数
            mod_params = mod.get("parameters", [])
            if mod_params:
                add_heading_cn(self.doc, f"5.{mod_num}.1 修饰参数", level=3)
                param_rows = []
                for param in mod_params:
                    param_rows.append([param.get("name", ""), param.get("value", "")])
                create_table(self.doc, ["参数", "设置"], param_rows, col_widths=[6, 10])

            # 修饰步骤
            mod_steps = mod.get("steps", [])
            if mod_steps:
                add_heading_cn(self.doc, f"5.{mod_num}.2 修饰步骤", level=3)
                for step in mod_steps:
                    step_num = step.get("step_number", "")
                    step_desc = step.get("description", "")
                    add_formatted_paragraph(self.doc, f"{step_num}. {step_desc}",
                                           font_size=FONT_SIZE_XIAOSI,
                                           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                           first_line_indent=True, space_before=1, space_after=1)

            add_page_break(self.doc)

    # ==================== 质量控制 ====================

    def _build_quality_control(self, data: Dict[str, Any]):
        """构建质量控制方案"""
        add_heading_cn(self.doc, "六、质量控制方案", level=1)
        add_horizontal_line(self.doc)

        qc = data.get("quality_control", {})

        # 分析策略
        add_heading_cn(self.doc, "6.1 分析策略", level=2)
        strategy = qc.get("strategy", "")
        if strategy:
            add_formatted_paragraph(self.doc, strategy,
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=6, space_after=6)

        # 检测项目
        add_heading_cn(self.doc, "6.2 检测项目", level=2)
        tests = qc.get("tests", [])
        test_headers = ["检测项目", "方法", "标准", "目的"]
        test_rows = []
        for test in tests:
            test_rows.append([
                test.get("name", ""),
                test.get("method", ""),
                test.get("standard", ""),
                test.get("purpose", "")
            ])
        
        if test_rows:
            create_table(self.doc, test_headers, test_rows, col_widths=[4, 4, 4, 4])

        # 验收标准
        add_heading_cn(self.doc, "6.3 验收标准", level=2)
        standards = qc.get("acceptance_standards", [])
        for standard in standards:
            add_formatted_paragraph(self.doc, f"● {standard}",
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=2, space_after=2)

        add_page_break(self.doc)

    # ==================== 项目周期与报价 ====================

    def _build_timeline_and_pricing(self, data: Dict[str, Any]):
        """构建项目周期与报价"""
        add_heading_cn(self.doc, "七、项目周期与报价", level=1)
        add_horizontal_line(self.doc)

        timeline = data.get("timeline", {})
        pricing = data.get("pricing", {})

        # 项目周期
        add_heading_cn(self.doc, "7.1 项目周期", level=2)
        phases = timeline.get("phases", [])
        phase_headers = ["阶段", "工作内容", "周期(工作日)", "交付物"]
        phase_rows = []
        for phase in phases:
            phase_rows.append([
                phase.get("name", ""),
                phase.get("tasks", ""),
                str(phase.get("duration", "")),
                phase.get("deliverables", "")
            ])
        
        if phase_rows:
            create_table(self.doc, phase_headers, phase_rows, col_widths=[3, 6, 3, 4])

        total_duration = timeline.get("total_duration", "")
        if total_duration:
            add_formatted_paragraph(self.doc, f"总周期：{total_duration}",
                                   font_size=FONT_SIZE_XIAOSI, bold=True,
                                   alignment=WD_ALIGN_PARAGRAPH.LEFT,
                                   first_line_indent=False, space_before=6, space_after=6)

        # 报价明细
        add_heading_cn(self.doc, "7.2 报价明细", level=2)
        items = pricing.get("items", [])
        price_headers = ["项目", "单价", "数量", "金额", "备注"]
        price_rows = []
        for item in items:
            price_rows.append([
                item.get("name", ""),
                item.get("unit_price", ""),
                str(item.get("quantity", "")),
                item.get("amount", ""),
                item.get("notes", "")
            ])
        
        if price_rows:
            create_table(self.doc, price_headers, price_rows, col_widths=[5, 3, 2, 3, 3])

        total_price = pricing.get("total", "")
        if total_price:
            add_formatted_paragraph(self.doc, f"合计：{total_price}",
                                   font_size=FONT_SIZE_SIHAO, bold=True,
                                   alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                                   first_line_indent=False, space_before=6, space_after=6)

        # 付款方式
        payment = pricing.get("payment_terms", "")
        if payment:
            add_formatted_paragraph(self.doc, f"付款方式：{payment}",
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=4, space_after=4)

        add_page_break(self.doc)

    # ==================== 风险评估与建议 ====================

    def _build_risk_assessment(self, data: Dict[str, Any]):
        """构建风险评估与建议"""
        add_heading_cn(self.doc, "八、风险评估与建议", level=1)
        add_horizontal_line(self.doc)

        risks = data.get("risks", [])
        
        if risks:
            risk_headers = ["风险类别", "风险描述", "影响程度", "应对措施"]
            risk_rows = []
            for risk in risks:
                risk_rows.append([
                    risk.get("category", ""),
                    risk.get("description", ""),
                    risk.get("impact", ""),
                    risk.get("mitigation", "")
                ])
            create_table(self.doc, risk_headers, risk_rows, col_widths=[3, 6, 2, 5])

        # 专业建议
        add_heading_cn(self.doc, "8.1 专业建议", level=2)
        suggestions = data.get("suggestions", [])
        for suggestion in suggestions:
            add_formatted_paragraph(self.doc, f"● {suggestion}",
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=2, space_after=2)

        add_page_break(self.doc)

    # ==================== 附录 ====================

    def _build_appendix(self, data: Dict[str, Any]):
        """构建附录"""
        add_heading_cn(self.doc, "附录", level=1)
        add_horizontal_line(self.doc)

        appendix = data.get("appendix", {})

        # 术语表
        add_heading_cn(self.doc, "附录A  术语表", level=2)
        glossary = appendix.get("glossary", [])
        if glossary:
            glossary_headers = ["术语/缩写", "中文全称", "定义"]
            glossary_rows = [[g.get("term", ""), g.get("term_cn", ""), g.get("definition", "")]
                            for g in glossary]
            create_table(self.doc, glossary_headers, glossary_rows, col_widths=[3, 4, 9])

        # 参考文献
        add_heading_cn(self.doc, "附录B  参考文献", level=2)
        refs = appendix.get("references", [])
        for ref in refs:
            num = ref.get("number", "")
            authors = ref.get("authors", "")
            title = ref.get("title", "")
            journal = ref.get("journal", "")
            year = ref.get("year", "")
            
            ref_text = f"[{num}] {authors} {title}. {journal}, {year}."
            add_formatted_paragraph(self.doc, ref_text,
                                   font_size=FONT_SIZE_WUHAO,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=False, space_before=2, space_after=2)

        # 免责声明
        add_heading_cn(self.doc, "附录C  免责声明", level=2)
        declarations = [
            "1. 本方案基于客户提供的需求信息编制，仅供技术参考。",
            "2. 实际合成结果可能受多种因素影响，具体以实验数据为准。",
            "3. 未经书面许可，不得将本方案用于商业目的或向第三方披露。",
            "4. 本方案的最终解释权归多肽AI辅助设计中心所有。",
        ]

        for decl in declarations:
            add_formatted_paragraph(self.doc, decl,
                                   font_size=FONT_SIZE_XIAOSI,
                                   alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                                   first_line_indent=True, space_before=4, space_after=4)

        # 签章区域
        for _ in range(4):
            self.doc.add_paragraph()

        add_mixed_paragraph(self.doc, [
            {"text": "方案编制人签字：", "bold": True, "font_size": FONT_SIZE_SIHAO,
             "font_cn": FONT_TITLE_CN, "font_en": FONT_TITLE_EN},
            {"text": "_________________", "font_size": FONT_SIZE_SIHAO}
        ], alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
           space_before=6, space_after=12)

        add_mixed_paragraph(self.doc, [
            {"text": "审核人签字：", "bold": True, "font_size": FONT_SIZE_SIHAO,
             "font_cn": FONT_TITLE_CN, "font_en": FONT_TITLE_EN},
            {"text": "_________________", "font_size": FONT_SIZE_SIHAO}
        ], alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
           space_before=6, space_after=12)

        add_mixed_paragraph(self.doc, [
            {"text": "日期：", "bold": True, "font_size": FONT_SIZE_SIHAO,
             "font_cn": FONT_TITLE_CN, "font_en": FONT_TITLE_EN},
            {"text": "_________________", "font_size": FONT_SIZE_SIHAO}
        ], alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=False,
           space_before=6, space_after=12)
